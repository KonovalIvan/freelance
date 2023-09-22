import os
import threading
from tkinter import filedialog, messagebox
import tkinter as tk
from tkinter.ttk import Entry, Progressbar
from typing import Optional
from xml.etree.ElementTree import Element, SubElement, tostring
import xml.dom.minidom

import docx
from docx.opc.exceptions import PackageNotFoundError

from exceptions import ParsingStoppedError, ParsingPathError


class XmlParserServices:

    def __init__(self):
        self.start_button = None
        self.dir_text_field = None
        self.stop_button = None
        self.thread_parsing = None
        self.thread_progress = None
        self.event = threading.Event()

    @staticmethod
    def choose_directory(dir_text_field: Optional[Entry]):
        directory = filedialog.askopenfilename()
        if dir_text_field:
            dir_text_field.delete(0, tk.END)
            dir_text_field.insert(0, directory)
        return directory

    def parse_xml(self, dir_text_field: tk.Entry, progress_bar: Progressbar):
        try:
            doc = docx.Document(dir_text_field.get())
        except PackageNotFoundError:
            messagebox.showinfo("Information", "The path contains an error.")
            self.finish_parsing_xml(self.stop_button, self.start_button, self.dir_text_field, False)
            raise ParsingPathError
        root = Element("document")

        # Var for progress bar
        total_characters = sum(len(paragraph.text) for paragraph in doc.paragraphs)
        processed_characters = 0

        # Logic for parsing pattern
        for paragraph in doc.paragraphs:
            if self.event.is_set():
                raise ParsingStoppedError
            paragraf_element = SubElement(root, "paragraph")
            paragraf_element.text = paragraph.text

            processed_characters += len(paragraph.text)
            progress = (processed_characters / total_characters) * 100
            self.progres_bar(progress_bar, progress)

        # File saving
        xml_string = xml.dom.minidom.parseString(tostring(root)).toprettyxml()
        xml_file_path = os.path.splitext(dir_text_field.get())[0] + ".xml"

        with open(xml_file_path, "w", encoding="utf-8") as xml_file:
            xml_file.write(xml_string)

        messagebox.showinfo("Information", "Parsing completed successfully")
        self.finish_parsing_xml(self.stop_button, self.start_button, self.dir_text_field, False)
        return True

    def progres_bar(self, progress_bar: Progressbar, progress: float):
        progress_bar["value"] = progress
        progress_bar.update_idletasks()

    def finish_parsing_xml(
            self, stop_button: Optional[tk.Button],
            start_button: Optional[tk.Button],
            dir_text_field: Optional[tk.Entry],
            stop: bool
    ):
        if stop:
            answer = messagebox.askyesno("Confirmation", "Are you sure you want to stop parsing?")
            if answer:
                self.event.set()
                messagebox.showinfo("Information", "Parsing was stop.")
        stop_button.config(state=tk.DISABLED) if stop_button else None
        start_button.config(state=tk.NORMAL) if start_button else None
        dir_text_field.config(state=tk.NORMAL) if dir_text_field else None

    def start_parsing_xml(self, stop_button: tk.Button, start_button: tk.Button, dir_text_field: tk.Entry,
                          progress_bar: Progressbar):
        self.stop_button = stop_button
        self.start_button = start_button
        self.dir_text_field = dir_text_field

        stop_button.config(state=tk.NORMAL)
        start_button.config(state=tk.DISABLED)
        dir_text_field.config(state=tk.DISABLED)
        self.thread_parsing = threading.Thread(target=self.parse_xml,
                                               args=(dir_text_field, progress_bar))
        self.thread_parsing.start()
